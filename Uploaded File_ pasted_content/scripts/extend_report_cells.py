from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches
ROOT=Path('/home/ubuntu/indic-peft-comparison'); p=ROOT/'deliverables'/'final_capstone_report.docx'; doc=Document(p); df=pd.read_csv(ROOT/'results/05-full-experiment-sweep/experiment_results.csv')
doc.add_page_break(); doc.add_heading('APPENDIX H — CELL-WISE ANALYTICAL RECORD',1)
for lang,lname in [('hi','Hindi'),('te','Telugu')]:
 for b in [50,100,500,1000,2000,20000]:
  doc.add_page_break(); d=df[(df.language==lang)&(df.budget==b)]
  doc.add_heading(f'{lname} — {b:,}-example cell',2)
  g=d.groupby('method').macro_f1.agg(['mean','std']).sort_values('mean',ascending=False)
  order=' > '.join(g.index.map({'lora':'LoRA','dora':'DoRA','ia3':'IA³'}))
  q=doc.add_paragraph(f'The observed primary ordering in this cell is {order}. The means and standard deviations are: '+', '.join(f'{m.upper()} mean={r["mean"]:.4f}, SD={r["std"]:.4f}' for m,r in g.iterrows())+'.')
  q.paragraph_format.first_line_indent=Inches(.35); q.paragraph_format.line_spacing=1.5
  q=doc.add_paragraph('This cell is interpreted jointly with its three seed outputs, the near-baseline diagnostic, and the confidence interval. The ordering is a descriptive result of the recorded protocol. It does not establish a statistically significant difference, and it should not be extrapolated beyond the tested language, model, task, or budget. The purpose of retaining this record is to make every ranking decision auditable from the underlying runs.')
  q.paragraph_format.first_line_indent=Inches(.35); q.paragraph_format.line_spacing=1.5
  t=doc.add_table(rows=1,cols=4); t.style='Table Grid'
  for j,x in enumerate(['Method','Seed 42','Seed 123','Seed 456']): t.rows[0].cells[j].text=x
  for m in ['lora','dora','ia3']:
   cells=t.add_row().cells; cells[0].text={'lora':'LoRA','dora':'DoRA','ia3':'IA³'}[m]
   vals=d[d.method==m].sort_values('seed').macro_f1.tolist()
   for j,v in enumerate(vals,1): cells[j].text=f'{v:.4f}'
doc.save(p); print('cell appendix added')
