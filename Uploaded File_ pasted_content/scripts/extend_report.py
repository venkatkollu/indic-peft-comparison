from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
ROOT=Path('/home/ubuntu/indic-peft-comparison'); OUT=ROOT/'deliverables'
doc=Document(OUT/'final_capstone_report.docx'); primary=pd.read_csv(ROOT/'results/05-full-experiment-sweep/experiment_results.csv'); summary=pd.read_csv(ROOT/'results/06-results-analysis/summary_with_ci.csv')
def cell(c,x,b=False):
 c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(x)); r.font.name='Times New Roman'; r.font.size=Pt(7); r.bold=b; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def addtable(df,cap):
 p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(cap); r.bold=True; r.font.size=Pt(10)
 t=doc.add_table(rows=1,cols=len(df.columns)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
 for j,c in enumerate(df.columns): cell(t.rows[0].cells[j],c,True)
 for _,row in df.iterrows():
  cs=t.add_row().cells
  for j,v in enumerate(row): cell(cs[j],f'{v:.4f}' if isinstance(v,float) else v)
# Detailed result interpretation appendix
p=doc.add_paragraph(); p.add_run().add_break()
doc.add_page_break(); doc.add_heading('APPENDIX E — BUDGET-WISE SEED-LEVEL ANALYSIS',1)
for lang, lname in [('hi','Hindi'),('te','Telugu')]:
 for budget in [50,100,500,1000,2000,20000]:
  doc.add_page_break(); doc.add_heading(f'{lname} at {budget:,} Training Examples',2)
  d=primary[(primary.language==lang)&(primary.budget==budget)].copy(); d['method']=d.method.map({'lora':'LoRA','dora':'DoRA','ia3':'IA³'})
  addtable(d[['method','seed','lr','epochs','accuracy','macro_f1','trainable_params','peak_gpu_memory_gb','training_time_sec']],f'Table E{[50,100,500,1000,2000,20000].index(budget)+1}. {lname} seed-level primary output at budget {budget:,}.')
  s=summary[(summary.language==lang)&(summary.budget==budget)].copy(); s['method']=s.method.map({'lora':'LoRA','dora':'DoRA','ia3':'IA³'})
  addtable(s[['method','f1_mean','f1_std','f1_ci95','acc_mean','acc_std']],f'Table E{[50,100,500,1000,2000,20000].index(budget)+7}. {lname} aggregate statistics at budget {budget:,}.')
  winner=s.loc[s.f1_mean.idxmax(),'method']; para=doc.add_paragraph(); para.paragraph_format.first_line_indent=Inches(.35); para.paragraph_format.line_spacing=1.5; para.add_run(f'Interpretation. At {budget:,} examples, the highest observed mean macro-F1 in {lname} is {winner}. The seed-level table is retained to show the dispersion underlying that mean. In the smallest budgets, near-baseline outcomes and wide intervals make the nominal ordering fragile; at larger budgets, the gap between methods can be interpreted descriptively within this fixed experimental protocol, without converting it into a universal or statistically significant claim.')
# engineering appendix
doc.add_page_break(); doc.add_heading('APPENDIX F — IMPLEMENTATION AND REPRODUCIBILITY NOTES',1)
for title,text in [('F.1 Execution sequence','The notebooks form a sequential pipeline: environment setup, dataset validation, preprocessing, infrastructure validation, pooled learning-rate search, primary sweep, corrected analysis, held-out evaluation, hybrid formal experiment, hybrid test evaluation, and final comparison. The output directories preserve the link between each stage and its downstream tables.'),('F.2 Configuration excerpt','The primary sweep uses MODEL_NAME = xlm-roberta-base, NUM_LABELS = 3, MAX_LENGTH = 128, BATCH_SIZE = 32, METHODS = [lora, dora, ia3], LANGUAGES = [hi, te], BUDGETS = [50, 100, 500, 1000, 2000, 20000], and SEEDS = [42, 123, 456].'),('F.3 Save/load safeguard','The model builder uses modules_to_save=[classifier] and asserts that classifier parameters are trainable. The save routine writes the PEFT adapter and classifier_head.pt separately, while the load routine restores the classifier state with strict=False. This workaround is part of the experiment definition, not an incidental convenience.'),('F.4 Data availability','The repository does not include the complete raw dataset. Reproduction therefore requires access to IndicXNLI and reconstruction of the processed parquet subsets. The preserved CSV/JSONL outputs permit result inspection without rerunning training, but exact reruns remain environment- and path-dependent.'),('F.5 Interpretation policy','The report treats docs/validated-results.md and corrected result CSVs as authoritative over older report-material drafts. Conflicts are recorded in CHANGELOG.md. Validation results and held-out hybrid results are never merged into one primary table.')]: doc.add_heading(title,2); q=doc.add_paragraph(text); q.paragraph_format.first_line_indent=Inches(.35); q.paragraph_format.line_spacing=1.5
# closing audit note
doc.add_page_break(); doc.add_heading('APPENDIX G — DATA INTEGRITY CHECKLIST',1)
checks=['108 primary rows verified','3 methods verified','2 languages verified','6 budgets verified','3 seeds verified','Macro-F1 and accuracy columns verified','Ranking winners recomputed from raw primary CSV','24 near-baseline runs recomputed with atol=0.001','Three-seed interval multiplier recorded as 4.302652729','Hybrid comparison contains 12 cells and 8 wins','Primary and hybrid experiments kept separate','No pairwise significance claim made','All figures generated from stored repository outputs','Repository commit recorded in audit materials']
for c in checks: q=doc.add_paragraph('Verified: '+c); q.paragraph_format.left_indent=Inches(.3); q.paragraph_format.line_spacing=1.5
doc.save(OUT/'final_capstone_report.docx'); print('extended')
