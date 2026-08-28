from pathlib import Path
from docx import Document
p=Path('/home/ubuntu/indic-peft-comparison/deliverables/final_capstone_report.docx'); out=p.parent/'final_capstone_report_source.md'; d=Document(p); lines=[]
for para in d.paragraphs:
 txt=para.text.strip()
 if not txt: continue
 style=para.style.name if para.style else ''
 if style.startswith('Heading 1'): lines += ['# '+txt,'']
 elif style.startswith('Heading 2'): lines += ['## '+txt,'']
 elif style.startswith('Heading 3'): lines += ['### '+txt,'']
 else: lines += [txt,'']
for ti,t in enumerate(d.tables,1):
 if not t.rows: continue
 lines.append(f'\n**Table {ti}.**\n')
 rows=[[c.text.replace('\n',' ').strip() for c in r.cells] for r in t.rows]
 lines.append('| '+' | '.join(rows[0])+' |'); lines.append('| '+' | '.join(['---']*len(rows[0]))+' |')
 for row in rows[1:]: lines.append('| '+' | '.join(row)+' |')
 lines.append('')
out.write_text('\n'.join(lines)); print(out)
